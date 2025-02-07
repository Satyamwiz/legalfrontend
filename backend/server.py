from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
# from flask_limiter import Limiter
# from flask_limiter.util import get_remote_address

import os
import time
import traceback
import hashlib
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor

from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, TextLoader, CSVLoader, UnstructuredExcelLoader
from langchain.schema import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from docx import Document as DocxDocument
from pypdf import PdfReader

app = Flask(__name__)
CORS(app)

# Uncomment below if you wish to enable rate limiting in production
# limiter = Limiter(get_remote_address, app=app, default_limits=["100 per hour", "10 per minute"])

load_dotenv()

groq_api_key = os.getenv('GROQ_API_KEY')
google_api_key = os.getenv("GOOGLE_API_KEY")
if not groq_api_key or not google_api_key:
    raise ValueError("Missing required API keys in .env file")

app.config["GOOGLE_API_KEY"] = google_api_key

# Initialize LLM with desired parameters
llm = ChatGroq(
    groq_api_key=groq_api_key, 
    model_name="Llama3-8b-8192",
    temperature=0.7,
    max_tokens=4096
)

# Global storage for processed files and vectors
processed_files = {}
vectors = None
executor = ThreadPoolExecutor(max_workers=4)

UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20MB

# Prompts
LEGAL_EXTRACTION_PROMPT = """You are a precise legal document analyzer. Follow these steps to extract ALL information, including every date, timeline, and related context:

1. First thoroughly analyze the entire document context
2. Extract information in these detailed categories:

ALL DATES AND TIME-RELATED INFORMATION
- List every single date mentioned in the document
- For each date, provide:
  * The exact date
  * Its context (what the date refers to)
  * Related conditions or requirements
  * Any associated deadlines or milestones
- Look for date-related terms like:
  * "as of", "effective", "commencing", "starting"
  * "until", "through", "ending", "terminating"
  * "within", "by", "no later than"
  * "renewal", "extension", "expiration"
  * Any specific day, month, or year mentioned

PARTIES AND RELATIONSHIPS
- List every entity or person mentioned
- For each party, include:
  * Their full name/designation
  * Their role in the document
  * Any responsibilities or obligations
  * Relationships with other parties
  * Associated terms or conditions

FINANCIAL DETAILS
- Extract all monetary amounts
- For each financial term:
  * The amount and currency
  * Purpose of payment
  * Payment schedule or timeline
  * Related conditions
  * Associated dates or deadlines

OBLIGATIONS AND REQUIREMENTS
- List all requirements found
- For each obligation:
  * Who is responsible
  * What needs to be done
  * When it needs to be done
  * Related conditions
  * Consequences of non-compliance

LEGAL AND COMPLIANCE
- All legal terms mentioned
- Each compliance requirement
- Regulatory references
- Governing laws
- Jurisdictional details

SPECIAL TERMS
- Identify any unique or special provisions
- Note unusual requirements
- Flag critical conditions
- Highlight important limitations

Format each section with:
1. Main heading
2. Clear bullet points
3. Context for each point
4. Related cross-references

For each extracted item:
- Include the surrounding context
- Note any dependencies or relationships
- Highlight critical implications
- Flag any ambiguities or unclear terms

If any information appears in multiple contexts, list each occurrence with its specific context.
If information is unclear or not specified, state "Not explicitly specified, but related context suggests [explanation]"

Context for analysis:
{context}

Please provide your detailed extraction:"""

SUMMARY_PROMPT = """You are an expert legal document summarizer. Create a concise and clear summary of the provided document following these guidelines:

DOCUMENT OVERVIEW
- Document type and purpose
- Key parties involved
- Overall scope

CORE AGREEMENT TERMS
- Main rights and obligations
- Critical deadlines and dates
- Financial arrangements

KEY LEGAL POINTS
- Governing law and major compliance requirements

NOTABLE PROVISIONS
- Special clauses and limitations

PRACTICAL IMPLICATIONS
- Main takeaways and action items

Ensure your summary is:
- Brief and to the point (around 100–150 words)
- Well-structured with clear headings
- Focused on essential details

Context to summarize:
{context}

Please provide your concise summary:"""

QA_PROMPT = """You are a legal document expert assistant. Analyze the following context and question carefully:

Context:
{context}

Question:
{input}

Provide a comprehensive answer that:
1. Directly addresses the question
2. Cites specific sections/clauses when relevant
3. Explains any legal implications
4. Highlights related terms or conditions
5. Notes any ambiguities or potential concerns

If the information isn't explicitly stated in the context, indicate that and explain any standard legal practices that might apply."""
    
extraction_prompt = ChatPromptTemplate.from_template(
    f"{LEGAL_EXTRACTION_PROMPT}\n\nDocument Content:\n{{context}}\n\nExtracted Information:"
)

qa_prompt = ChatPromptTemplate.from_template(QA_PROMPT)

# File handling functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def is_pdf_text_extractable(file_path):
    try:
        reader = PdfReader(file_path)
        text_content = ""
        for page in reader.pages[:2]:
            text_content += page.extract_text()
        return len(text_content.strip()) > 100
    except Exception as e:
        print(f"PDF verification error: {e}")
        return False

def process_pdf(file_path):
    if is_pdf_text_extractable(file_path):
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            if not any(len(doc.page_content.strip()) > 100 for doc in documents):
                raise ValueError("Extracted content appears to be insufficient")
            return documents
        except Exception as e:
            print(f"PDF processing error: {e}")
            raise
    else:
        raise ValueError("PDF appears to be image-based or corrupted. Text extraction not supported.")

def process_uploaded_file(file):
    try:
        file_name = file.filename
        file_extension = file_name.rsplit('.', 1)[1].lower()
        
        if not allowed_file(file_name):
            raise ValueError(f"File type {file_extension} not supported")
        
        tmp_path = os.path.join(UPLOAD_FOLDER, file_name)
        file.save(tmp_path)

        with open(tmp_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()

        if file_hash in processed_files:
            print(f"File {file_name} already processed, skipping...")
            return []

        if file_extension == 'pdf':
            documents = process_pdf(tmp_path)
        elif file_extension == 'docx':
            doc = DocxDocument(tmp_path)
            documents = [Document(page_content=para.text) for para in doc.paragraphs if para.text.strip()]
        elif file_extension == 'txt':
            documents = TextLoader(tmp_path).load()
        elif file_extension == 'csv':
            documents = CSVLoader(tmp_path).load()
        elif file_extension == 'xlsx':
            documents = UnstructuredExcelLoader(tmp_path).load()
        else:
            return []

        if not documents or not any(len(doc.page_content.strip()) > 50 for doc in documents):
            raise ValueError(f"No valid content extracted from {file_name}")

        processed_files[file_hash] = documents
        return documents

    except Exception as e:
        print(f"❌ Error processing {file_name}: {e}\n{traceback.format_exc()}")
        raise

# Vectorization with chunking
def vector_embedding(documents):
    global vectors
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=250,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
        is_separator_regex=False
    )
    
    final_documents = text_splitter.split_documents(documents)
    
    try:
        if vectors is None:
            vectors = FAISS.from_documents(final_documents, embeddings)
        else:
            vectors.add_documents(final_documents)
    except Exception as e:
        print(f"Embedding error: {e}")
        raise

@app.route('/')
def index():
    return render_template('index.html')

# Upload endpoint: processes files, creates document chunks, and submits for vectorization
@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400

    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({"error": "No selected files"}), 400

    all_documents = []
    processed_file_names = []
    errors = []
    extracted_text = ""

    for file in files:
        try:
            if file and allowed_file(file.filename):
                file.seek(0)
                file_content = file.read()
                file_hash = hashlib.md5(file_content).hexdigest()
                file.seek(0)
                
                if file_hash not in processed_files:
                    docs = process_uploaded_file(file)
                    all_documents.extend(docs)
                    processed_files[file_hash] = docs
                    processed_file_names.append(file.filename)
                else:
                    print(f"File {file.filename} already processed, skipping...")
        except Exception as e:
            error_msg = f"Error processing {file.filename}: {str(e)}"
            print(f"❌ {error_msg}\n{traceback.format_exc()}")
            errors.append(error_msg)

    if all_documents:
        try:
            executor.submit(vector_embedding, all_documents)
            extracted_text = "\n\n".join([
                f"From {processed_file_names[i]}: {doc.page_content[:1000]}..."
                for i, doc in enumerate(all_documents[:3])
            ])
        except Exception as e:
            errors.append(f"Error in vector embedding: {str(e)}")

    response = {
        "message": "Files processed successfully ✅" if not errors else "Partial success with errors",
        "processed_files": processed_file_names,
        "extracted_text": extracted_text,
        "status": "success" if not errors else "partial_success",
        "processed_count": len(processed_file_names)
    }
    if errors:
        response["errors"] = errors

    return jsonify(response)

# Extraction endpoint: returns detailed extraction using the full extraction prompt
@app.route('/extract', methods=['GET'])
def extract_details():
    if not vectors:
        return jsonify({"error": "No documents uploaded"}), 400

    try:
        document_chain = create_stuff_documents_chain(llm, extraction_prompt)
        retriever = vectors.as_retriever(
            search_type="similarity",
            search_kwargs={
                'k': 8,
                'fetch_k': 20,
                'lambda_mult': 0.8
            }
        )
        retrieval_chain = create_retrieval_chain(retriever, document_chain)

        start = time.process_time()
        response = retrieval_chain.invoke({
            'input': 'Perform a comprehensive analysis and extract all legal details from all document sections'
        })
        elapsed = time.process_time() - start

        return jsonify({
            "answer": response['answer'],
            "time_taken": f"{elapsed:.2f}s",
            "context_chunks": len(response.get('context', [])),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "error": f"Extraction failed: {str(e)}",
            "status": "error"
        }), 500

# Ask endpoint: processes a legal question based on the uploaded documents
@app.route('/ask', methods=['POST'])
def ask_question():
    if not vectors:
        return jsonify({"error": "No documents uploaded"}), 400

    data = request.json
    user_question = data.get("query")

    if not user_question:
        return jsonify({"error": "No question provided"}), 400

    try:
        document_chain = create_stuff_documents_chain(llm, qa_prompt)
        retriever = vectors.as_retriever(
            search_type="similarity",
            search_kwargs={
                'k': 8,
                'fetch_k': 20,
                'lambda_mult': 0.8
            }
        )
        retrieval_chain = create_retrieval_chain(retriever, document_chain)

        start = time.process_time()
        response = retrieval_chain.invoke({"input": user_question})
        elapsed = time.process_time() - start

        return jsonify({
            "answer": response['answer'],
            "time_taken": f"{elapsed:.2f}s",
            "context_chunks": len(response.get('context', [])),
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "error": f"Question processing failed: {str(e)}",
            "status": "error"
        }), 500

# Summary endpoint: returns a concise summary using reduced retrieval parameters to save tokens
@app.route('/summary', methods=['POST'])
def ask_summary():
    if not vectors:
        return jsonify({"error": "No documents uploaded"}), 400

    try:
        # Create prompt template using the concise summary prompt
        summary_prompt_template = ChatPromptTemplate.from_template(SUMMARY_PROMPT)
        
        # Create document chain with the updated prompt
        document_chain = create_stuff_documents_chain(
            llm=llm,
            prompt=summary_prompt_template,
        )

        # Use reduced retrieval parameters to lower token usage
        retriever = vectors.as_retriever(
            search_type="similarity",
            search_kwargs={
                'k': 5,           # Fewer, more focused chunks
                'fetch_k': 15,    # Limit token consumption
                'lambda_mult': 0.7,
                'filter': None
            }
        )

        # Create and invoke chain for a single concise summary
        retrieval_chain = create_retrieval_chain(retriever, document_chain)
        start = time.process_time()
        response = retrieval_chain.invoke({
            'input': 'Provide a concise summary of the document focusing on key points'
        })
        elapsed = time.process_time() - start

        return jsonify({
            "answer": response['answer'],
            "time_taken": f"{elapsed:.2f}s",
            "context_chunks": len(response.get('context', [])),
            "status": "success"
        })

    except Exception as e:
        print(f"Summary error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            "error": f"Summary generation failed: {str(e)}",
            "status": "error"
        }), 500
        

if __name__ == '__main__':
    app.run(debug=False)
